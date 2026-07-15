import csv
import io
from dataclasses import dataclass
from pathlib import PurePath

from app.core.errors import AppError

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".xlsx", ".csv"}
MAX_ATTACHMENT_BYTES = 25 * 1024 * 1024
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 160


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    warning: str | None = None
    page_count: int | None = None


def validate_filename(filename: str | None) -> str:
    safe_name = PurePath(filename or "").name
    suffix = PurePath(safe_name).suffix.lower()
    if not safe_name or safe_name in {".", ".."} or suffix not in SUPPORTED_EXTENSIONS:
        raise AppError(
            code="unsupported_attachment_type",
            message="Only txt, md, docx, pdf, xlsx and csv files are supported.",
            status_code=400,
            details={"supported_extensions": sorted(SUPPORTED_EXTENSIONS)},
        )
    return safe_name


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    validate_filename(filename)
    if len(content) > MAX_ATTACHMENT_BYTES:
        raise AppError(
            code="attachment_too_large",
            message="The attachment exceeds the 25 MB limit.",
            status_code=413,
        )

    suffix = PurePath(filename).suffix.lower()
    try:
        if suffix in {".txt", ".md"}:
            return ParsedDocument(_decode_text(content))
        if suffix == ".csv":
            return ParsedDocument(_parse_csv(content))
        if suffix == ".xlsx":
            return ParsedDocument(_parse_xlsx(content))
        if suffix == ".docx":
            return ParsedDocument(_parse_docx(content))
        return _parse_pdf(content)
    except AppError:
        raise
    except Exception as error:  # noqa: BLE001 - parser errors become readable API state
        raise AppError(
            code="attachment_parse_failed",
            message="The attachment could not be parsed.",
            status_code=422,
            details={"reason": str(error)},
        ) from error


def split_into_chunks(text: str) -> list[str]:
    normalized = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n")).strip()
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + CHUNK_SIZE, len(normalized))
        if end < len(normalized):
            boundary = normalized.rfind("\n", start + CHUNK_SIZE // 2, end)
            if boundary > start:
                end = boundary
        chunks.append(normalized[start:end].strip())
        if end >= len(normalized):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)
    return [chunk for chunk in chunks if chunk]


def _decode_text(content: bytes) -> str:
    return content.decode("utf-8-sig", errors="replace")


def _parse_csv(content: bytes) -> str:
    decoded = _decode_text(content)
    rows = csv.reader(io.StringIO(decoded))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _parse_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sections: list[str] = []
    for sheet in workbook.worksheets:
        sections.append(f"[工作表: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                sections.append(" | ".join(values))
    return "\n".join(sections)


def _parse_docx(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    sections = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            sections.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(sections)


def _parse_pdf(content: bytes) -> ParsedDocument:
    import fitz

    document = fitz.open(stream=content, filetype="pdf")
    pages = [page.get_text("text").strip() for page in document]
    text = "\n\n".join(page for page in pages if page)
    if not text:
        return ParsedDocument(
            text="",
            warning="该 PDF 未提取到文本，可能是扫描版 PDF。请上传可复制文本的 PDF 或 OCR 后再试。",
            page_count=document.page_count,
        )
    warning = None
    if any(not page for page in pages):
        warning = "该 PDF 部分页面未提取到文本，可能包含扫描页；已索引可读取页面。"
    return ParsedDocument(text=text, warning=warning, page_count=document.page_count)
